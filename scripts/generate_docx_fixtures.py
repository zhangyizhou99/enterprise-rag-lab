"""Generate reproducible DOCX fixtures for parser and cleaner validation.

Run with:
    d:/Code/hello-agents/.venv/Scripts/python.exe scripts/generate_docx_fixtures.py
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FASTAPI_ROOT = PROJECT_ROOT / "data" / "raw" / "fastapi"
RAW_DOCX_DIR = PROJECT_ROOT / "data" / "raw" / "docx"
FIXTURE_DIR = PROJECT_ROOT / "data" / "fixtures" / "docx"
FASTAPI_COMMIT = "afe41126f624af30038cc8e17b2aaf60ebd4b838"


def _set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)


def _add_header_and_footer(document: Document, label: str) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"Enterprise RAG Lab | {label}"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.text = "Synthetic test fixture. Not an internal company document."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_title(document: Document, title: str, subtitle: str) -> None:
    document.add_heading(title, level=0)
    paragraph = document.add_paragraph(subtitle)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_deployment_runbook() -> Path:
    document = Document()
    _set_default_font(document)
    _add_header_and_footer(document, "API Service Deployment Runbook")
    _add_title(
        document,
        "API Service Deployment Runbook",
        "Synthetic fixture for DOCX structure extraction",
    )

    document.add_heading("1. Purpose", level=1)
    document.add_paragraph(
        "This runbook describes the repeatable deployment process for a FastAPI service. "
        "It is deliberately concise and contains headings, lists, tables, and code blocks."
    )

    document.add_heading("2. Prerequisites", level=1)
    for item in [
        "A reviewed release tag.",
        "A validated environment configuration.",
        "Access to the deployment target and application logs.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("3. Environment variables", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["Name", "Required", "Purpose"]):
        cell.text = value
    for row in [
        ("APP_ENV", "Yes", "Selects the runtime environment."),
        ("DATABASE_URL", "Yes", "Connection string for the application database."),
        ("LOG_LEVEL", "No", "Controls structured application log verbosity."),
    ]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value

    document.add_heading("4. Deployment steps", level=1)
    for item in [
        "Run the automated test suite and inspect the release notes.",
        "Apply schema migrations before switching application traffic.",
        "Deploy the service, then issue a health check request.",
        "Monitor error rate and request latency for fifteen minutes.",
    ]:
        document.add_paragraph(item, style="List Number")

    document.add_heading("5. Health check command", level=1)
    command = document.add_paragraph()
    command.style = document.styles["Normal"]
    command.add_run("curl -fsS https://api.example.test/healthz").font.name = "Consolas"

    document.add_heading("6. Rollback", level=1)
    document.add_paragraph(
        "Rollback when the error rate exceeds the release threshold or the health check fails. "
        "Restore the last known-good release, verify the health endpoint, and record the incident."
    )

    path = FIXTURE_DIR / "api-service-deployment-runbook.docx"
    document.save(path)
    return path


def build_noisy_postmortem() -> Path:
    document = Document()
    _set_default_font(document)
    _add_header_and_footer(document, "CORS Incident Postmortem - Draft")
    _add_title(
        document,
        "CORS Incident Postmortem",
        "Synthetic noisy fixture for DOCX cleaning validation",
    )

    document.add_paragraph("DRAFT - DO NOT DISTRIBUTE", style="Subtitle")
    document.add_paragraph("DRAFT - DO NOT DISTRIBUTE", style="Subtitle")
    document.add_paragraph("")

    document.add_heading("1. Summary", level=1)
    document.add_paragraph(
        "Browser clients could not call the public API after a CORS policy update. "
        "The API returned a successful preflight response, but the response omitted the required origin header."
    )

    document.add_heading("2. Timeline", level=1)
    timeline = document.add_table(rows=1, cols=2)
    timeline.style = "Table Grid"
    timeline.rows[0].cells[0].text = "Time"
    timeline.rows[0].cells[1].text = "Event"
    for row in [
        ("09:10", "Deployment completed."),
        ("09:18", "Browser clients reported blocked cross-origin requests."),
        ("09:31", "Configuration mismatch was identified."),
        ("09:45", "A corrected allow-origin rule was deployed."),
    ]:
        cells = timeline.add_row().cells
        cells[0].text, cells[1].text = row

    document.add_heading("3. Root cause", level=1)
    document.add_paragraph(
        "The configured allow-origin value did not match the browser application's origin. "
        "The service also enabled credentials, which prohibits a wildcard origin."
    )

    document.add_heading("4. Resolution", level=1)
    document.add_paragraph(
        "The team replaced the wildcard configuration with an explicit allowlist and added an automated preflight test."
    )

    document.add_heading("5. Deliberate noise for cleaner tests", level=1)
    for text in [
        "Download the desktop client from https://example.invalid/download.",
        "Your browser does not support embedded media.",
        "DRAFT - DO NOT DISTRIBUTE",
        "DRAFT - DO NOT DISTRIBUTE",
    ]:
        document.add_paragraph(text)

    document.add_paragraph("")
    document.add_paragraph("End of synthetic fixture.")

    path = FIXTURE_DIR / "cors-incident-postmortem-noisy.docx"
    document.save(path)
    return path


def _add_markdown_line(document: Document, line: str, in_code_block: bool) -> None:
    if in_code_block:
        paragraph = document.add_paragraph()
        paragraph.add_run(line).font.name = "Consolas"
        return

    heading = re.match(r"^(#{1,6})\s+(.+?)(?:\s+\{[^}]+\})?$", line)
    if heading:
        document.add_heading(heading.group(2), level=len(heading.group(1)))
    elif line.startswith("* "):
        document.add_paragraph(line[2:], style="List Bullet")
    elif re.match(r"^\d+\.\s+", line):
        document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
    elif line and not line.startswith(("{*", "///", "<img ", "![")):
        document.add_paragraph(line)


def build_official_fastapi_copy(source_relative: str) -> Path:
    source_path = FASTAPI_ROOT / source_relative
    lines = source_path.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].split(" { #", 1)[0] for line in lines if line.startswith("# ")), source_path.stem)

    document = Document()
    _set_default_font(document)
    _add_header_and_footer(document, "Official FastAPI Documentation Copy")
    _add_title(document, title, "Official FastAPI Chinese documentation - local format copy")
    document.add_paragraph(
        "Source: https://github.com/fastapi/fastapi/blob/"
        f"{FASTAPI_COMMIT}/docs/zh/{source_relative.replace('\\', '/')}"
    )
    document.add_paragraph(
        "This file is an attributed local DOCX representation of public FastAPI documentation. "
        "It is used only for local parsing, retrieval evaluation, and demonstration."
    )

    in_code_block = False
    for line in lines:
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue
        if line.startswith("# "):
            continue
        _add_markdown_line(document, line, in_code_block)

    path = RAW_DOCX_DIR / f"fastapi-{source_path.stem}.docx"
    document.save(path)
    return path


def main() -> None:
    RAW_DOCX_DIR.mkdir(parents=True, exist_ok=True)
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    source_documents = [
        "docs/tutorial/cors.md",
        "docs/tutorial/dependencies/classes-as-dependencies.md",
        "docs/tutorial/dependencies/dependencies-with-yield.md",
        "docs/tutorial/dependencies/global-dependencies.md",
        "docs/advanced/advanced-dependencies.md",
    ]
    paths = [build_deployment_runbook(), build_noisy_postmortem()]
    paths.extend(build_official_fastapi_copy(source_document) for source_document in source_documents)
    for path in paths:
        print(path.relative_to(PROJECT_ROOT))


if __name__ == "__main__":
    main()
