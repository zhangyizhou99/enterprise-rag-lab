"""DOCX parser for paragraphs, tables, headers, and footers."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document as OpenDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from enterprise_rag_lab.models import ParsedBlock, ParseResult, SourceFormat

_HEADING_STYLE = re.compile(r"Heading\s+(\d+)", re.IGNORECASE)


class DocxParser:
    def parse(self, path: Path) -> ParseResult:
        document = OpenDocument(path)
        blocks: list[ParsedBlock] = []
        heading_stack: list[str] = []
        title = path.stem
        table_index = 0

        def append_block(
            text: str,
            block_type: str,
            metadata: dict[str, object] | None = None,
        ) -> None:
            normalized = text.strip()
            if not normalized:
                return
            blocks.append(
                ParsedBlock(
                    ordinal=len(blocks),
                    text=normalized,
                    block_type=block_type,
                    heading_path=tuple(heading_stack),
                    metadata=metadata or {},
                )
            )

        for item in document.iter_inner_content():
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if not text:
                    continue
                style_name = item.style.name if item.style else ""
                heading_match = _HEADING_STYLE.fullmatch(style_name)
                if style_name.lower() == "title":
                    title = text
                    append_block(text, "title", {"style": style_name})
                elif heading_match:
                    level = int(heading_match.group(1))
                    heading_stack[level - 1 :] = [text]
                    append_block(text, "heading", {"level": level, "style": style_name})
                else:
                    append_block(text, "paragraph", {"style": style_name})
            elif isinstance(item, Table):
                rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in item.rows]
                append_block(
                    "\n".join(row for row in rows if row.strip()),
                    "table",
                    {"table_index": table_index},
                )
                table_index += 1

        seen_headers_and_footers: set[tuple[str, str]] = set()
        for section_index, section in enumerate(document.sections):
            for block_type, paragraphs in (
                ("header", section.header.paragraphs),
                ("footer", section.footer.paragraphs),
            ):
                text = "\n".join(paragraph.text.strip() for paragraph in paragraphs if paragraph.text.strip())
                identity = (block_type, text)
                if text and identity not in seen_headers_and_footers:
                    append_block(text, block_type, {"section_index": section_index})
                    seen_headers_and_footers.add(identity)

        return ParseResult(
            source_path=path,
            source_format=SourceFormat.DOCX,
            title=title,
            text="\n\n".join(block.text for block in blocks),
            blocks=tuple(blocks),
            metadata={
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "section_count": len(document.sections),
            },
        )